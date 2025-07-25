# index_builder.py

import os
import json
import argparse
from pathlib import Path
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.schema import Document
from PyPDF2 import PdfReader
import docx
from dotenv import load_dotenv

# --- Load environment variable for OpenAI API key (local GPT-3.5 setup) ---
load_dotenv("GPT35.env")

# --- Config ---
INDEX_DIR = "embeddings"
BOOK_SOURCES = {
    "Australian Slang": {
        "type": "json",
        "path": "data/slang_book/book_pdf_pages_labeled.json",
        "index": "slang_index"
    },
    "Periodic Table": {
        "type": "json",
        "path": "data/periodic_book/book_pdf_pages_labeled.json",
        "index": "periodic_index"
    },
    "MySQL Workshop": {
        "type": "json",
        "path": "data/mysql_book/book_pdf_pages_labeled.json",
        "index": "mysql_index"
    },
    "Excel Budgeting": {
        "type": "json",
        "path": "data/excel_book/book_pdf_pages_labeled.json",
        "index": "excel_index"
    }
}

# --- Helpers ---
def build_json_index(source):
    with open(source, "r", encoding="utf-8") as f:
        data = json.load(f)

    docs = []
    for entry in data:
        metadata = {
            "page": entry.get("page"),
            "pdf_page": entry.get("pdf_page"),
            "pg_safe": True
        }
        docs.append(Document(page_content=entry.get("text", ""), metadata=metadata))
    return docs

def build_pdf_index(source):
    reader = PdfReader(source)
    docs = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        metadata = {"page": i + 1, "pg_safe": True}
        docs.append(Document(page_content=text or "", metadata=metadata))
    return docs

def build_docx_index(source):
    doc = docx.Document(source)
    docs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            metadata = {"entry_id": i + 1, "pg_safe": True}
            docs.append(Document(page_content=text, metadata=metadata))
    return docs

def build_index(name, config):
    print(f"\n🔧 Building index for: {name}")
    source_type = config["type"]
    source_path = config["path"]
    index_name = config["index"]

    if not os.path.exists(source_path):
        print(f"❌ Source not found: {source_path}")
        return

    if source_type == "json":
        docs = build_json_index(source_path)
    elif source_type == "pdf":
        docs = build_pdf_index(source_path)
    elif source_type == "docx":
        docs = build_docx_index(source_path)
    else:
        print(f"❌ Unknown source type: {source_type}")
        return

    # Inject short-term memory from session state if available
    if os.getenv("BAZZA_CONTEXT"):
        docs.insert(0, Document(page_content=os.getenv("BAZZA_CONTEXT"), metadata={"pg_safe": True}))

    vectorstore = FAISS.from_documents(docs, OpenAIEmbeddings(openai_api_key=os.getenv("OPENAI_API_KEY")))
    save_path = os.path.join(INDEX_DIR, index_name)
    Path(save_path).mkdir(parents=True, exist_ok=True)
    vectorstore.save_local(save_path)
    print(f"✅ Saved to {save_path}")

# --- Main ---
def main():
    for book, config in BOOK_SOURCES.items():
        build_index(book, config)

if __name__ == "__main__":
    main()

